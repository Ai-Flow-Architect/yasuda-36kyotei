"""
全7様式 PDF/Word 生成精度テスト
Excelから読み込み → PDF+Word生成 → 内容検証 → 精度スコア出力
"""
import sys
import os
import difflib
import re
import traceback
from pathlib import Path

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from excel_reader import read_excel
from pdf_generator import generate_pdf
from word_generator import generate_word, FORM_NAMES

# ── 依存ライブラリ（オプション: PDF テキスト抽出）
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

OUTPUT_DIR = "test_output_all7"
EXCEL_PATH = "demo_data/test_all7forms.xlsx"
ERRORS = []
RESULTS = []


def _extract_pdf_text(pdf_bytes: bytes) -> str:
    """PDFからテキスト抽出（PyMuPDF使用）"""
    if not HAS_PYMUPDF:
        return ""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    return "\n".join(page.get_text() for page in doc)


def _check_pdf_content(text: str, record: dict, pat: str) -> dict:
    """PDFテキストの内容チェック"""
    scores = {}
    company = record.get("事業所名", "").replace("（様式9）", "").replace("（様式9_2）", "")\
        .replace("（様式9_3）", "").replace("（様式9_4）", "").replace("（様式9_5）", "")\
        .replace("（様式10）", "").replace("（様式10_2）", "")

    # ── 基本フィールドチェック
    checks = {
        "事業所名": company in text if company else True,
        "代表者名": record.get("事業主名", "") in text,
        "労働者代表": record.get("労働者代表_氏名", "") in text,
        "起算日": f"令和{record.get('起算日_年', '')}年" in text,
        "延長時間_月": record.get("延長時間_1ヶ月", "") in text if record.get("延長時間_1ヶ月") else True,
        "協定の文字列": "協定" in text,
    }

    # ── 様式別チェック
    # 様式9_2/10/10_2: 特別条項の月時間・回数を確認
    # 様式9_3/9_4/9_5: これらの様式はシンプルテーブル形式のため特別条項詳細を別テーブルに持たない → チェックしない
    if pat in ("9_2", "10", "10_2"):
        checks["特別条項_月時間"] = record.get("特別_延長時間_月", "") in text if record.get("特別_延長時間_月") else True
        checks["特別条項_超過回数"] = record.get("特別_超過回数", "") in text if record.get("特別_超過回数") else True

    if pat in ("10", "10_2"):
        checks["1年変形制"] = "1年" in text and "変形" in text
        checks["第2章"] = "第2章" in text

    passed = sum(1 for v in checks.values() if v)
    total = len(checks)
    score = int(passed / total * 100) if total else 0

    return {
        "checks": checks,
        "passed": passed,
        "total": total,
        "score": score,
    }


def _extract_word_text(doc) -> str:
    """Wordから全テキストを再帰的に抽出（ネストされたテーブル含む）"""
    parts = []
    # 段落
    for p in doc.paragraphs:
        parts.append(p.text)
    # テーブル（ネストも含む）
    def extract_table(table):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)
                for nested_tbl in cell.tables:
                    extract_table(nested_tbl)
    for table in doc.tables:
        extract_table(table)
    return "\n".join(parts)


def _check_word_content(docx_path: str, record: dict, pat: str) -> dict:
    """Wordファイルの内容チェック"""
    if not HAS_DOCX:
        return {"score": 0, "error": "python-docxなし"}

    doc = DocxDocument(docx_path)
    text = _extract_word_text(doc)
    return _check_pdf_content(text, record, pat)


def run_test():
    print("=" * 70)
    print("全7様式 生成精度テスト")
    print("=" * 70)

    # Excelからデータ読み込み
    if not Path(EXCEL_PATH).exists():
        print(f"❌ Excelファイルが見つかりません: {EXCEL_PATH}")
        sys.exit(1)

    records = read_excel(EXCEL_PATH)
    print(f"\n📊 読み込み: {len(records)} 件")
    for r in records:
        print(f"   - {r['事業所名']} → 様式{r['様式パターン']}")

    Path(OUTPUT_DIR).mkdir(exist_ok=True)

    print("\n" + "─" * 70)
    print("PDFおよびWord生成テスト")
    print("─" * 70)

    for record in records:
        pat = record.get("様式パターン", "9")
        name = record.get("事業所名", "不明")
        form_label = FORM_NAMES.get(pat, f"様式{pat}")

        print(f"\n▶ {name} | 様式{pat} ({form_label})")

        result = {
            "事業所名": name,
            "様式パターン": pat,
            "PDF生成": False,
            "Word生成": False,
            "PDF精度スコア": 0,
            "Word精度スコア": 0,
            "エラー": [],
        }

        # ─── PDF生成
        try:
            pdf_bytes = generate_pdf(record)
            safe_name = re.sub(r'[\\/:*?"<>|（）]', '_', name)
            pdf_path = Path(OUTPUT_DIR) / f"36協定書_{safe_name}_{pat}.pdf"
            pdf_path.write_bytes(pdf_bytes)
            result["PDF生成"] = True
            print(f"  ✅ PDF生成: {pdf_path.name} ({len(pdf_bytes):,} bytes)")

            # テキスト抽出して精度チェック
            if HAS_PYMUPDF:
                pdf_text = _extract_pdf_text(pdf_bytes)
                check_result = _check_pdf_content(pdf_text, record, pat)
                result["PDF精度スコア"] = check_result["score"]
                print(f"  📊 PDF内容チェック: {check_result['passed']}/{check_result['total']} 項目合格 ({check_result['score']}%)")
                for k, v in check_result["checks"].items():
                    status = "✅" if v else "❌"
                    print(f"       {status} {k}")
            else:
                print("  ⚠️  PyMuPDFなし → テキスト抽出スキップ")

        except Exception as e:
            result["エラー"].append(f"PDF: {e}")
            ERRORS.append((name, pat, "PDF生成", str(e)))
            print(f"  ❌ PDF生成エラー: {e}")
            if "--verbose" in sys.argv:
                traceback.print_exc()

        # ─── Word生成
        try:
            word_path = generate_word(record, OUTPUT_DIR)
            result["Word生成"] = True
            size = Path(word_path).stat().st_size
            print(f"  ✅ Word生成: {Path(word_path).name} ({size:,} bytes)")

            # Wordの内容チェック
            if HAS_DOCX:
                check_result = _check_word_content(word_path, record, pat)
                result["Word精度スコア"] = check_result["score"]
                if "error" not in check_result:
                    print(f"  📊 Word内容チェック: {check_result['passed']}/{check_result['total']} 項目合格 ({check_result['score']}%)")

        except Exception as e:
            result["エラー"].append(f"Word: {e}")
            ERRORS.append((name, pat, "Word生成", str(e)))
            print(f"  ❌ Word生成エラー: {e}")
            if "--verbose" in sys.argv:
                traceback.print_exc()

        RESULTS.append(result)

    # ─── 結果サマリー
    print("\n" + "=" * 70)
    print("テスト結果サマリー")
    print("=" * 70)
    print(f"\n{'様式':<8} {'事業所名':<28} {'PDF':>5} {'Word':>5} {'PDF精度':>8} {'Word精度':>8}")
    print("─" * 70)

    total_pdf_score = 0
    total_word_score = 0
    pdf_ok = 0
    word_ok = 0

    for r in RESULTS:
        pat = r["様式パターン"]
        name = r["事業所名"][:26]
        pdf_status = "✅" if r["PDF生成"] else "❌"
        word_status = "✅" if r["Word生成"] else "❌"
        pdf_score = r["PDF精度スコア"]
        word_score = r["Word精度スコア"]

        if r["PDF生成"]:
            pdf_ok += 1
            total_pdf_score += pdf_score
        if r["Word生成"]:
            word_ok += 1
            total_word_score += word_score

        print(f"  {pat:<6}  {name:<28} {pdf_status:>5} {word_status:>5} {pdf_score:>7}% {word_score:>7}%")

    total = len(RESULTS)
    print("─" * 70)
    avg_pdf = int(total_pdf_score / pdf_ok) if pdf_ok else 0
    avg_word = int(total_word_score / word_ok) if word_ok else 0
    print(f"  合計({total}件)  PDF成功:{pdf_ok}/{total}  Word成功:{word_ok}/{total}")
    if HAS_PYMUPDF:
        print(f"  PDF内容精度 平均: {avg_pdf}%")
    if HAS_DOCX:
        print(f"  Word内容精度 平均: {avg_word}%")

    if ERRORS:
        print(f"\n❌ エラー {len(ERRORS)} 件:")
        for name, pat, kind, msg in ERRORS:
            print(f"   [{pat}] {name}: {kind} → {msg}")
        sys.exit(1)
    else:
        print(f"\n✅ 全{total}様式 生成成功！")
        if HAS_PYMUPDF and avg_pdf < 99:
            print(f"⚠️  PDF内容精度 {avg_pdf}% (目標: 99%)")
        sys.exit(0)


if __name__ == "__main__":
    run_test()
