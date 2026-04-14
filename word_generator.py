"""
36協定書 Word生成モジュール
python-docxで協定書を動的生成する（7様式パターン対応）
"""
import logging
import re
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from pathlib import Path

# ---------------------------------------------------------------------------
# ロガー設定
# ---------------------------------------------------------------------------
logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# 定数定義
# ---------------------------------------------------------------------------
# フォント
DEFAULT_FONT_NAME: str = "游明朝"

# フォントサイズ (pt)
FONT_SIZE_TITLE: float = 14
FONT_SIZE_SUBTITLE: float = 11
FONT_SIZE_BODY: float = 10.5
FONT_SIZE_ARTICLE: float = 10
FONT_SIZE_TABLE_CELL: float = 9
FONT_SIZE_SEPARATOR: float = 8
FONT_SIZE_NOTE: float = 9
FONT_SIZE_SPECIAL_TITLE: float = 12
FONT_SIZE_CHAPTER: float = 12

# 段落の後方スペース (デフォルト)
DEFAULT_SPACE_AFTER: Pt = Pt(6)
TITLE_SPACE_AFTER: Pt = Pt(12)

# ページ設定 (mm / cm)
PAGE_WIDTH_MM: int = 210
PAGE_HEIGHT_MM: int = 297
MARGIN_CM: float = 2.5

# テーブルスタイル
TABLE_STYLE: str = "Table Grid"

# 区切り線
SEPARATOR_CHAR: str = "─"
SEPARATOR_LENGTH: int = 40
DOUBLE_SEPARATOR_CHAR: str = "="
DOUBLE_SEPARATOR_LENGTH: int = 60


# ---------------------------------------------------------------------------
# ユーティリティ関数
# ---------------------------------------------------------------------------
def _v(data: Dict[str, Any], key: str, default: str = "") -> str:
    """辞書から値を取得してstr変換。Noneや空文字の場合はdefaultを返す。"""
    val = data.get(key)
    if val is None:
        return default
    s = str(val).strip()
    return s if s else default


def _safe_filename(name: str) -> str:
    """ファイル名に使用できない文字（/ \\ : * ? " < > | & % #）をアンダースコアに置換"""
    name = name.replace(" ", "").replace("　", "")
    name = re.sub(r'[\\/:*?"<>|&%#]', "_", name)
    name = name.strip(". ")
    return name or "不明"


def set_cell_text(
    cell: _Cell,
    text: str,
    bold: bool = False,
    size: float = FONT_SIZE_ARTICLE,
    align: int = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    """セルにテキストを設定"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = DEFAULT_FONT_NAME
    run.bold = bold


def add_paragraph(
    doc: Document,
    text: str,
    bold: bool = False,
    size: float = FONT_SIZE_BODY,
    align: int = WD_ALIGN_PARAGRAPH.LEFT,
    space_after: Pt = DEFAULT_SPACE_AFTER,
) -> Paragraph:
    """段落を追加"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = space_after
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = DEFAULT_FONT_NAME
    run.bold = bold
    return p


def _add_table(
    doc: Document,
    rows_data: List[Tuple[str, str]],
) -> None:
    """ラベル-値ペアのリストから2列テーブルを追加する"""
    table = doc.add_table(rows=len(rows_data), cols=2, style=TABLE_STYLE)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(rows_data):
        set_cell_text(table.cell(i, 0), label, bold=True, size=FONT_SIZE_TABLE_CELL)
        set_cell_text(table.cell(i, 1), value, size=FONT_SIZE_TABLE_CELL)


def _start_date(data: Dict[str, Any]) -> str:
    """起算日を「令和○年○月1日」形式で返す。未入力の場合は「〇」プレースホルダ。"""
    y = _v(data, "起算日_年") or "〇"
    m = _v(data, "起算日_月") or "〇"
    return f"令和{y}年{m}月1日"


def _health_measure(data: Dict[str, Any]) -> str:
    """特別_健康措置の3段階フォールバック（PDF側と統一）"""
    内容 = _v(data, "特別_健康措置_内容")
    番号 = _v(data, "特別_健康措置_番号")
    if 内容:
        return 内容
    if 番号:
        return f"番号{番号}の措置"
    return "医師による面接指導"


# ---------------------------------------------------------------------------
# 共通パーツ: ヘッダー（タイトル + 事業所情報 + 前文）
# ---------------------------------------------------------------------------
def _generate_header(
    doc: Document,
    data: Dict[str, Any],
    *,
    title: Optional[str] = None,
    subtitle: Optional[str] = None,
    preamble: Optional[str] = None,
) -> None:
    """全様式共通のヘッダー部分を生成する

    Args:
        doc: Wordドキュメント
        data: 入力データ辞書
        title: 協定書タイトル（Noneの場合は標準タイトル）
        subtitle: タイトル直下に表示するサブタイトル（様式9_3〜9_5用）
        preamble: 前文テキスト（Noneの場合は様式9号用のデフォルト前文を使用）
    """
    # タイトル
    doc_title = title or "時間外労働及び休日労働に関する協定書"
    add_paragraph(
        doc,
        doc_title,
        bold=True,
        size=FONT_SIZE_TITLE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=TITLE_SPACE_AFTER,
    )

    # サブタイトル（様式9_3〜9_5で使用）
    if subtitle:
        add_paragraph(
            doc, subtitle, size=FONT_SIZE_SUBTITLE, align=WD_ALIGN_PARAGRAPH.CENTER
        )

    # 事業所情報
    add_paragraph(doc, f"事業の名称: {_v(data, '事業所名')}")
    add_paragraph(doc, f"事業の種類: {_v(data, '事業の種類')}")
    if subtitle is None and preamble is None:
        # 様式9号（標準）は電話番号付き
        add_paragraph(
            doc,
            f"事業主名: {_v(data, '事業主名')}　（電話: {_v(data, '電話番号')}）",
        )
    else:
        add_paragraph(doc, f"事業主名: {_v(data, '事業主名')}")

    # 前文
    add_paragraph(doc, "")
    if preamble is not None:
        add_paragraph(doc, preamble, size=FONT_SIZE_ARTICLE)
    else:
        # 様式9号用デフォルト前文（PDF側と社名フィールドを統一: 事業所名を使用）
        add_paragraph(
            doc,
            f"{_v(data, '事業所名')}（以下「甲」という。）と労働者代表者（以下「乙」という。）は、"
            "労働基準法第３６条第１項の規定に基づき、労働基準法に定める法定労働時間"
            "（１週４０時間、１日８時間）並びに変形労働時間制に定める所定労働時間を超えた労働時間で、"
            "かつ１日８時間、１週４０時間の法定労働時間又は変形期間の法定労働時間の総枠を超える労働"
            "（以下「時間外労働」という。）及び労働基準法に定める休日（毎週１日又は４週４日）における労働"
            "（以下「休日労働」という。）に関し、次の通り協定する。",
            size=FONT_SIZE_ARTICLE,
        )


# ---------------------------------------------------------------------------
# 共通パーツ: 署名欄（フッター）
# ---------------------------------------------------------------------------
def _generate_footer(doc: Document, data: Dict[str, Any]) -> None:
    """全様式共通の署名欄を生成する"""
    add_paragraph(doc, "")
    add_paragraph(doc, SEPARATOR_CHAR * SEPARATOR_LENGTH, size=FONT_SIZE_SEPARATOR)
    add_paragraph(doc, "")

    add_paragraph(doc, f"協定締結日: {_v(data, '協定締結日')}")
    add_paragraph(doc, f"届出作成日: {_v(data, '届出作成日')}")
    add_paragraph(doc, "")
    add_paragraph(
        doc,
        f"（甲）{_v(data, '事業所名')}　{_v(data, '事業主職名', '代表取締役')}　{_v(data, '事業主名')}",
    )
    add_paragraph(
        doc,
        f"（乙）労働者代表　{_v(data, '労働者代表_職')}　{_v(data, '労働者代表_氏名')}",
    )
    add_paragraph(doc, "")
    add_paragraph(doc, f"所轄労働局: {_v(data, '所轄労働局')}")
    add_paragraph(doc, f"所轄労基署: {_v(data, '所轄労基署')}")


# 後方互換性のためエイリアスを残す
_add_signature_section = _generate_footer


# ---------------------------------------------------------------------------
# 共通パーツ: 簡易時間外労働テーブル（様式9_3/9_4/9_5共通）
# ---------------------------------------------------------------------------
def _generate_simple_overtime_table(
    doc: Document,
    data: Dict[str, Any],
    extra_rows: Optional[List[Tuple[str, str]]] = None,
) -> None:
    """様式9_3/9_4/9_5で共通する時間外労働テーブルを生成する"""
    add_paragraph(doc, "【時間外労働】", bold=True)
    rows: List[Tuple[str, str]] = [
        ("業務の種類", _v(data, "時間外_業務の種類")),
        ("労働者数", f"{_v(data, '労働者数')}人"),
        ("延長することができる時間（1日）", f"{_v(data, '延長時間_1日')}時間"),
        ("延長することができる時間（1ヶ月）", f"{_v(data, '延長時間_1ヶ月')}時間"),
    ]
    if extra_rows:
        rows.extend(extra_rows)
    _add_table(doc, rows)


# ---------------------------------------------------------------------------
# 共通パーツ: 1年変形制 第2章（様式10/10_2共通）
# ---------------------------------------------------------------------------
def _generate_chapter2_1nen(
    doc: Document,
    data: Dict[str, Any],
    art_start: int,
    起算日: str,
) -> None:
    """第2章「1年変形制について」を生成する（art_start = 最初の条番号）"""
    漢数字 = ["", "一", "二", "三", "四", "五", "六", "七", "八", "九",
              "一〇", "一一", "一二", "一三", "一四", "一五", "一六", "一七", "一八"]

    def art(no: int) -> str:
        return f"第{漢数字[no]}条" if no <= 18 else f"第{no}条"

    n = art_start
    所定時間 = _v(data, "所定労働時間", "〇時間〇分")

    # 始業終業時刻の分解
    始業終業_raw = _v(data, "始業終業時刻") or "〇時〇分　〇時〇分"
    if "〜" in 始業終業_raw:
        parts = 始業終業_raw.split("〜", 1)
        begin = parts[0].strip()
        end = parts[1].strip() if len(parts) > 1 else "〇時〇分"
    else:
        begin = 始業終業_raw
        end = "〇時〇分"
    rest = "〇時〇分より〇時〇分まで"

    add_paragraph(doc, "")
    add_paragraph(
        doc,
        "第2章　1年変形制について",
        bold=True,
        size=FONT_SIZE_CHAPTER,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(
        doc,
        "１年単位の変形労働時間制に関し、次のとおり協定する。",
        size=FONT_SIZE_ARTICLE,
    )
    add_paragraph(doc, "")

    # 第n条: 所定労働時間
    add_paragraph(
        doc,
        f"{art(n)}　所定労働時間は、1年単位の変形労働時間制によるものとし、"
        "1年を平均して週40時間を越えないものとする。",
        size=FONT_SIZE_ARTICLE,
    )
    add_paragraph(doc, "　１．　始業終業休憩時間は下記の通りとする。", size=FONT_SIZE_ARTICLE)
    _add_table(doc, [
        ("始業", begin),
        ("終業", end),
        ("休憩", rest),
    ])

    # 第n+1条: 対象者
    add_paragraph(doc, "")
    add_paragraph(
        doc,
        f"{art(n+1)}　本協定に基づく1年単位の変形労働時間制の対象者は次のとおりとする。",
        size=FONT_SIZE_ARTICLE,
    )
    add_paragraph(doc, "　１．　正規の当社社員。但し、次の者を除く。", size=FONT_SIZE_ARTICLE)
    add_paragraph(doc, "　　（１）　妊産婦で適用除外のあった者", size=FONT_SIZE_ARTICLE)
    add_paragraph(
        doc,
        "　　（２）　育児又は老人等の介護を行う者、職業訓練又は教育を受ける者、"
        "その他特別の理由により申し出をし、正当な事由より会社から認められた者",
        size=FONT_SIZE_ARTICLE,
    )
    add_paragraph(doc, "　　（３）　変形対象期間の中途で退職が予定される者及び中途採用者・配転者", size=FONT_SIZE_ARTICLE)
    add_paragraph(doc, "　２．　パートタイマー、アルバイト及び臨時の従業員", size=FONT_SIZE_ARTICLE)
    add_paragraph(doc, "　３．　嘱託雇用者で特に定めた者", size=FONT_SIZE_ARTICLE)

    # 第n+2条: 休日（原本準拠：「１週に１回の日曜日の休日を確保するものとし」）
    add_paragraph(doc, "")
    add_paragraph(
        doc,
        f"{art(n+2)}　前条の期間における休日は、１週に１回の日曜日の休日を確保するものとし、"
        "別添の年間カレンダーのとおりとする。尚、特定期間はないものとする。",
        size=FONT_SIZE_ARTICLE,
    )

    # 第n+3条: 所定労働日
    add_paragraph(
        doc,
        f"{art(n+3)}　前条の期間における所定労働日は、前条に定める休日以外の日とする。",
        size=FONT_SIZE_ARTICLE,
    )

    # 第n+4条: 変形対象期間の労働時間
    add_paragraph(
        doc,
        f"{art(n+4)}　変形対象期間における労働時間は、{所定時間}とする。（休憩時間を除く）",
        size=FONT_SIZE_ARTICLE,
    )

    # 第n+5条: 休日振替
    add_paragraph(
        doc,
        f"{art(n+5)}　業務上やむを得ない事由があるときは、前条の用件の範囲内で、"
        "従業員の代表の同意を得て同一週内において休日を振り替えることができる。",
        size=FONT_SIZE_ARTICLE,
    )

    # 第n+6条: 時間外手当
    add_paragraph(
        doc,
        f"{art(n+6)}　本協定に基づく所定労働時間を超えて労働した場合には、"
        "就業規則の規定に基づき時間外手当を支払うものとする。",
        size=FONT_SIZE_ARTICLE,
    )

    # 第n+7条: 有効期間
    add_paragraph(
        doc,
        f"{art(n+7)}　本協定の有効期間は、{起算日}から1年間とする。",
        size=FONT_SIZE_ARTICLE,
    )


# ---------------------------------------------------------------------------
# 各様式の固有部分
# ---------------------------------------------------------------------------
def _generate_form_9_articles_1to6(doc: Document, data: Dict[str, Any], 起算日: str) -> None:
    """様式9/9_2共通: 第1〜6条を生成するヘルパー（有効期間なし）"""
    # 第1条〜第3条
    add_paragraph(doc, "")
    add_paragraph(
        doc,
        "第１条　甲は、時間外労働及び休日労働を可能な限り行わせないように努める。",
        size=FONT_SIZE_ARTICLE,
    )
    add_paragraph(
        doc,
        "第２条　乙は、故意または過失により時間外労働及び休日労働を生じさせない義務を負う。",
        size=FONT_SIZE_ARTICLE,
    )
    add_paragraph(
        doc,
        "第３条　前2条にも関わらずその必要性を生じた場合、甲は次により時間外労働を行わせることができる。",
        size=FONT_SIZE_ARTICLE,
    )

    # 時間外労働テーブル
    add_paragraph(doc, "")
    _add_table(doc, [
        ("時間外労働をさせる必要のある具体的事由", _v(data, "時間外_事由")),
        ("業務の種類", _v(data, "時間外_業務の種類")),
        ("従事する労働者数（満18歳以上の者）", f"{_v(data, '労働者数')}人"),
        ("所定労働時間", f"{_v(data, '所定労働時間')}時間"),
        ("延長することができる時間（1日）", f"{_v(data, '延長時間_1日')}時間"),
        ("延長することができる時間（1ヶ月）", f"{_v(data, '延長時間_1ヶ月')}時間"),
        ("期間", _v(data, "時間外_期間", f"{起算日}から1年間")),
    ])

    # 第4条（時間外・休日の通知）
    add_paragraph(doc, "")
    add_paragraph(
        doc,
        "第４条　甲は、時間外労働を行わせる場合は、原則として、前日の終業時刻までに当該労働者に通知する。"
        "また、休日労働を行わせる場合は、原則として、前日の終業時刻までに当該労働者に通知する。",
        size=FONT_SIZE_ARTICLE,
    )

    # 第5条（起算日）― 原本A社・B社PDF準拠: 「第２条の表」
    add_paragraph(
        doc,
        f"第５条　第２条の表における1週、１ヶ月及び１年の起算日はいずれも{起算日}とする。",
        size=FONT_SIZE_ARTICLE,
    )

    # 第6条: 休日労働
    add_paragraph(
        doc,
        "第６条　甲は、必要がある場合には、次により休日労働を行わせることができる。",
        size=FONT_SIZE_ARTICLE,
    )
    add_paragraph(doc, "")
    _add_table(doc, [
        ("休日労働をさせる必要のある具体的事由", _v(data, "休日_事由")),
        ("業務の種類", _v(data, "休日_業務の種類")),
        ("所定休日", _v(data, "所定休日")),
        ("労働させることができる休日日数", f"1か月に{_v(data, '休日労働_日数')}日"),
        ("始業及び終業の時刻", _v(data, "始業終業時刻")),
    ])


def generate_form_9(doc: Document, data: Dict[str, Any]) -> None:
    """様式第9号: 一般条項（特別条項なし）― 原本A社PDF準拠・7条構成"""
    起算日 = _start_date(data)

    # ヘッダー（デフォルト前文）
    _generate_header(doc, data)

    # 第1〜6条
    _generate_form_9_articles_1to6(doc, data, 起算日)

    # 第7条: 有効期間（様式9は7条構成）
    add_paragraph(doc, "")
    add_paragraph(
        doc,
        f"第７条　本協定の有効期間は、{起算日}から１年間とする。",
        size=FONT_SIZE_ARTICLE,
    )

    # 署名欄
    _generate_footer(doc, data)


def generate_form_9_2(doc: Document, data: Dict[str, Any]) -> None:
    """様式第9号の2: 特別条項付き ― 原本B社PDF準拠・8条構成（第7条=段落形式特別条項）"""
    起算日 = _start_date(data)

    # ヘッダー（デフォルト前文）
    _generate_header(doc, data)

    # 第1〜6条（共通ヘルパー）
    _generate_form_9_articles_1to6(doc, data, 起算日)

    # デフォルト値をPDF側と統一
    理由   = _v(data, "特別_理由", "業務の繁忙")
    月時間  = _v(data, "特別_延長時間_月", "80")
    年時間  = _v(data, "特別_延長時間_年", "960")
    回数   = _v(data, "特別_超過回数", "6")
    割増率  = _v(data, "特別_割増賃金率", "25")
    手続き  = _v(data, "特別_手続き", "労使の協議")
    措置   = _health_measure(data)

    # 第7条: 特別条項（段落形式 ― PDFの_special_paraと同一文言）
    add_paragraph(doc, "")
    add_paragraph(
        doc,
        f"第７条　一定期間についての延長時間は1か月45時間とする。但し、{理由}に特に対応が必要な時は、"
        f"{手続き}を経て1か月{月時間}時間までこれを延長することができる。この場合、延長時間を更に"
        f"延長する回数と法定労働時間合計は、{回数}回及び{年時間}時間までとし、かつこれが1か月45時間"
        f"又は1年間360時間を超えた時間外労働に対しての割増率を{割増率}％とする。"
        f"また、60時間を超えた場合の割増率は50％とする。健康確保措置：{措置}。",
        size=FONT_SIZE_ARTICLE,
    )

    # 第8条: 有効期間
    add_paragraph(
        doc,
        f"第８条　本協定の有効期間は、{起算日}から１年間とする。",
        size=FONT_SIZE_ARTICLE,
    )

    # 署名欄
    _generate_footer(doc, data)


def generate_form_9_3(doc: Document, data: Dict[str, Any]) -> None:
    """様式第9号の3: 研究開発業務"""
    _generate_header(
        doc,
        data,
        subtitle="（新技術・新商品等の研究開発業務）",
        preamble=(
            "本協定は、労働基準法第36条第11項に基づき、新技術・新商品等の研究開発業務に"
            "従事する労働者に対し、時間外労働及び休日労働を行わせることについて協定する。"
        ),
    )
    add_paragraph(
        doc,
        "※ 研究開発業務は上限規制の適用除外ですが、月100時間超で医師の面接指導が義務付けられます。",
        size=FONT_SIZE_NOTE,
    )

    # 時間外労働テーブル
    _generate_simple_overtime_table(
        doc,
        data,
        extra_rows=[("期間", _v(data, "時間外_期間"))],
    )

    _generate_footer(doc, data)


def generate_form_9_4(doc: Document, data: Dict[str, Any]) -> None:
    """様式第9号の4: 適用猶予事業・業務"""
    _generate_header(
        doc,
        data,
        subtitle="（適用猶予期間中における適用猶予事業・業務）",
        preamble=(
            "本協定は、自動車運転者・建設業・医師等の適用猶予事業・業務に従事する"
            "労働者に対し、時間外労働及び休日労働を行わせることについて協定する。"
        ),
    )

    # 時間外労働テーブル
    _generate_simple_overtime_table(
        doc,
        data,
        extra_rows=[("期間", _v(data, "時間外_期間"))],
    )

    _generate_footer(doc, data)


def generate_form_9_5(doc: Document, data: Dict[str, Any]) -> None:
    """様式第9号の5: 適用猶予＋事業場外みなし労働"""
    _generate_header(
        doc,
        data,
        subtitle="（適用猶予事業・業務 / 事業場外みなし労働時間制）",
        preamble=(
            "本協定は、適用猶予事業・業務において、事業場外労働のみなし労働時間に係る協定の"
            "内容を36協定に付記して届出するものである。"
        ),
    )

    # 時間外労働テーブル（最終行が「事業場外みなし労働時間」になる点が他と異なる）
    _generate_simple_overtime_table(
        doc,
        data,
        extra_rows=[("事業場外みなし労働時間", f"{_v(data, '所定労働時間')}時間")],
    )

    _generate_footer(doc, data)


def generate_form_10(doc: Document, data: Dict[str, Any]) -> None:
    """様式第10号: 1年変形制（標準）"""
    起算日 = _start_date(data)

    # デフォルト値をPDF側と統一
    特別理由 = _v(data, "特別_理由", "業務繁忙")
    月時間   = _v(data, "特別_延長時間_月", "75")
    年時間   = _v(data, "特別_延長時間_年", "600")
    回数     = _v(data, "特別_超過回数", "6")
    limit    = _v(data, "延長時間_1ヶ月", "42")
    year_limit = "320"
    割増率   = _v(data, "特別_割増賃金率", "25")

    # ヘッダー
    _generate_header(
        doc,
        data,
        title="時間外労働及び休日労働及び1年変形制に関する協定書",
    )

    add_paragraph(doc, "")
    add_paragraph(
        doc,
        "第1章　時間外労働及び休日労働について",
        bold=True,
        size=FONT_SIZE_CHAPTER,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # 第1条〜第6条（様式9と共通）
    add_paragraph(
        doc,
        "第１条　甲は、時間外労働及び休日労働を可能な限り行わせないように努める。",
        size=FONT_SIZE_ARTICLE,
    )
    add_paragraph(
        doc,
        "第２条　乙は、故意または過失により時間外労働及び休日労働を生じさせない義務を負う。",
        size=FONT_SIZE_ARTICLE,
    )
    add_paragraph(
        doc,
        "第３条　前2条にも関わらずその必要性を生じた場合、甲は次により時間外労働を行わせることができる。",
        size=FONT_SIZE_ARTICLE,
    )

    add_paragraph(doc, "")
    add_paragraph(doc, "【時間外労働】", bold=True)
    _add_table(doc, [
        ("時間外労働をさせる必要のある具体的事由", _v(data, "時間外_事由")),
        ("業務の種類", _v(data, "時間外_業務の種類")),
        ("従事する労働者数（満18歳以上の者）", f"{_v(data, '労働者数')}人"),
        ("所定労働時間", f"{_v(data, '所定労働時間')}時間"),
        ("延長することができる時間（1日）", f"{_v(data, '延長時間_1日')}時間"),
        ("延長することができる時間（1ヶ月）", f"{_v(data, '延長時間_1ヶ月')}時間"),
    ])

    add_paragraph(doc, "")
    add_paragraph(
        doc,
        "第４条　甲は、時間外労働を行わせる場合は、原則として、前日の終業時刻までに当該労働者に通知する。"
        "また、休日労働を行わせる場合は、原則として、1日前の終業時刻までに当該労働者に通知する。",
        size=FONT_SIZE_ARTICLE,
    )
    # 第5条（起算日）― 様式10は「第３条の表」（原本D社PDF準拠）
    add_paragraph(
        doc,
        f"第５条　第３条の表における1週、１ヶ月及び１年の起算日はいずれも{起算日}とする。",
        size=FONT_SIZE_ARTICLE,
    )
    add_paragraph(
        doc,
        "第６条　甲は、必要がある場合には、次により休日労働を行わせることができる。",
        size=FONT_SIZE_ARTICLE,
    )

    add_paragraph(doc, "")
    _add_table(doc, [
        ("休日労働をさせる必要のある具体的事由", _v(data, "休日_事由")),
        ("業務の種類", _v(data, "休日_業務の種類")),
        ("所定休日", _v(data, "所定休日")),
        ("労働させることができる休日日数", f"1か月に{_v(data, '休日労働_日数')}日"),
        ("始業及び終業の時刻", _v(data, "始業終業時刻")),
    ])

    # 第7条（固定文言）
    add_paragraph(doc, "")
    add_paragraph(
        doc,
        "第７条　時間外労働及び休日労働の決定は会社によりなされ、よって会社の指示によってのみ行われる。"
        "つまり従業員の個人的な判断で行うことはできず、さらに包括的にその一切の権限をいかなる者にも与えることはない。",
        size=FONT_SIZE_ARTICLE,
    )

    # 第8条（特別条項）
    add_paragraph(
        doc,
        f"第８条　特別条項として{特別理由}に特に対応が必要なときは、労使の協議を経て、"
        f"１か月{月時間}時間までこれを延長することができる。この場合、延長時間を更に延長する回数と"
        f"法定労働外時間の合計は、{回数}回及び年間{年時間}時間までとし、かつこれが１ヶ月{limit}時間"
        f"又は１年間{year_limit}時間を超えた時間外労働に対しての割増率を{割増率}％とする。"
        "また、60時間を超える場合の割増率は50％とする。",
        size=FONT_SIZE_ARTICLE,
    )

    # 第2章（第9〜16条）
    _generate_chapter2_1nen(doc, data, art_start=9, 起算日=起算日)

    # 署名欄
    _generate_footer(doc, data)


def generate_form_10_2(doc: Document, data: Dict[str, Any]) -> None:
    """様式第10号の2: 1年変形制＋ドライバー（適用猶予）"""
    起算日 = _start_date(data)

    # デフォルト値をPDF側と統一
    割増率  = _v(data, "特別_割増賃金率", "25")
    手続き  = _v(data, "特別_手続き", "労働者代表者との協議による合意")
    措置   = _v(data, "特別_健康措置_内容") or "休日の確保 面接による健康の把握・労働状況の把握・改善指導"

    # ヘッダー
    _generate_header(
        doc,
        data,
        title="時間外労働及び休日労働及び1年変形制に関する協定書",
        subtitle="（適用猶予事業・業務 / 1年単位の変形労働時間制）",
        preamble=(
            f"{_v(data, '事業所名')}（以下「甲」という。）と労働者代表者（以下「乙」という。）は、"
            "労働基準法第３６条第１項の規定に基づき、労働基準法に定める法定労働時間"
            "（１週４０時間、１日８時間）並びに変形労働時間制に定める所定労働時間を超えた労働時間で、"
            "かつ１日８時間、１週４０時間の法定労働時間又は変形期間の法定労働時間の総枠を超える労働"
            "（以下「時間外労働」という。）及び労働基準法に定める休日（毎週１日又は４週４日）における労働"
            "（以下「休日労働」という。）に関し、次の通り協定する。"
        ),
    )

    add_paragraph(doc, "")
    add_paragraph(
        doc,
        "第1章　時間外労働及び休日労働について",
        bold=True,
        size=FONT_SIZE_CHAPTER,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # 第1条〜第2条（努力義務・乙の義務）
    add_paragraph(
        doc,
        "第１条　甲は、時間外労働及び休日労働を可能な限り行わせないように努める。",
        size=FONT_SIZE_ARTICLE,
    )
    add_paragraph(
        doc,
        "第２条　乙は、故意または過失により時間外労働及び休日労働を生じさせない義務を負う。",
        size=FONT_SIZE_ARTICLE,
    )

    # 第3条（時間外労働 + 改善基準告示）
    add_paragraph(
        doc,
        "第３条　前２条にも関わらずその必要性を生じた場合、甲は次により時間外労働を行わせることができる。",
        size=FONT_SIZE_ARTICLE,
    )
    add_paragraph(doc, "")
    add_paragraph(doc, "【時間外労働】", bold=True)
    _add_table(doc, [
        ("時間外労働をさせる必要のある具体的事由", _v(data, "時間外_事由")),
        ("業務の種類", _v(data, "時間外_業務の種類")),
        ("従事する労働者数（満18歳以上の者）", f"{_v(data, '労働者数')}人"),
        ("延長することができる時間（1日）", f"{_v(data, '延長時間_1日')}時間"),
        ("延長することができる時間（1ヶ月）", f"{_v(data, '延長時間_1ヶ月')}時間"),
        ("延長することができる時間（1年）", "360時間"),
    ])
    add_paragraph(
        doc,
        "　２　自動車運転者については、前項の規定により時間外労働を行わせることによって"
        "「自動車運転者の労働時間等の改善のための基準」（以下「改善基準告示」という。）に定める"
        "１か月及び１年についての拘束時間並びに１日についての最大拘束時間の限度を超えることとなる場合においては、"
        "当該拘束時間の限度をもって、前項の時間外労働時間の限度とする。",
        size=FONT_SIZE_ARTICLE,
    )

    # 第4条（休日労働 + 改善基準告示）
    add_paragraph(doc, "")
    add_paragraph(
        doc,
        "第４条　甲は、就業規則の規定に基づき、必要がある場合には、次により休日労働を行わせることができる。",
        size=FONT_SIZE_ARTICLE,
    )
    add_paragraph(doc, "")
    add_paragraph(doc, "【休日労働】", bold=True)
    _add_table(doc, [
        ("休日労働をさせる必要のある具体的事由", _v(data, "休日_事由")),
        ("業務の種類", _v(data, "休日_業務の種類")),
        ("所定休日", _v(data, "所定休日")),
        ("労働させることができる休日日数", f"1か月に{_v(data, '休日労働_日数')}日"),
        ("始業及び終業の時刻", _v(data, "始業終業時刻")),
    ])
    add_paragraph(
        doc,
        "　２　自動車運転者については、前項の規定により休日労働を行わせることによって、"
        "改善基準告示に定める１か月及び１年についての拘束時間並びに１日についての最大拘束時間の限度を"
        "超えることとなる場合においては、当該拘束時間の限度をもって、前項の休日労働の限度とする。",
        size=FONT_SIZE_ARTICLE,
    )

    # 第5条（特別条項）
    add_paragraph(doc, "")
    add_paragraph(
        doc,
        "第５条　通常予見することのできない業務量の大幅な増加等に伴う臨時的な場合であって、"
        "次の何れかに該当する場合は、第２条の規定に基づき時間外労働を行わせることができる時間を超えて労働させることができる。",
        size=FONT_SIZE_ARTICLE,
    )
    add_paragraph(doc, "")
    add_paragraph(doc, "【特別条項】", bold=True)
    _add_table(doc, [
        ("臨時的に限度時間を超える理由", _v(data, "特別_理由", "業務の繁忙")),
        ("業務の種類", _v(data, "特別_業務の種類")),
        ("労働者数", f"{_v(data, '特別_労働者数')}人"),
        ("限度時間を超えることができる回数", f"{_v(data, '特別_超過回数', '6')}回"),
        ("延長することができる時間数（月）", f"{_v(data, '特別_延長時間_月', '80')}時間"),
        ("延長することができる時間数（年）", f"{_v(data, '特別_延長時間_年', '960')}時間"),
    ])
    add_paragraph(
        doc,
        f"　２　前項の規定に基づいて限度時間を超えて労働させる場合の割増率は{割増率}％とする。"
        "　なお、時間外労働が１か月６０時間を超えた場合の割増率は５０％とする。",
        size=FONT_SIZE_ARTICLE,
    )
    add_paragraph(
        doc,
        "　３　第１項の規定に基づいて限度時間を超えて労働させる場合における手続及び"
        "限度時間を超えて労働させる労働者に対する健康及び福祉を確保するための措置については、次のとおりとする。",
        size=FONT_SIZE_ARTICLE,
    )
    add_paragraph(
        doc,
        f"　　限度時間を超えて労働させる場合における手続：{手続き}",
        size=FONT_SIZE_ARTICLE,
    )
    add_paragraph(
        doc,
        f"　　限度時間を超えて労働させる労働者に対する健康及び福祉を確保するための措置：{措置}",
        size=FONT_SIZE_ARTICLE,
    )
    add_paragraph(
        doc,
        "　４　自動車運転者については、第１項の規定により時間外労働を行わせることによって"
        "改善基準告示に定める１か月及び１年についての拘束時間並びに１日についての最大拘束時間の限度を"
        "超えることとなる場合においては、当該拘束時間の限度をもって、第１項の時間外労働の時間の限度とする。",
        size=FONT_SIZE_ARTICLE,
    )

    # 第6条（100時間・80時間平均規制）
    add_paragraph(doc, "")
    add_paragraph(
        doc,
        "第６条　第２条から第４条までの規定に基づいて時間外労働又は休日労働を行わせる場合においても、"
        "自動車運転者については、各条に定める時間数等にかかわらず、時間外労働及び休日労働を合算した時間数は"
        "１か月について１００時間未満となるよう努めるものとする。",
        size=FONT_SIZE_ARTICLE,
    )
    add_paragraph(
        doc,
        "　２　自動車運転者以外の者については、各条により定める時間数等にかかわらず、"
        "時間外労働及び休日労働を合算した時間数は、１か月について１００時間未満でなければならず、"
        "かつ２か月から６か月までを平均して８０時間を超過しないこととする。",
        size=FONT_SIZE_ARTICLE,
    )

    # 第7条（運転時間制限）
    add_paragraph(
        doc,
        "第７条　第２条から第４条までの規定に基づいて時間外労働又は休日労働を行わせる場合においても、"
        "自動車運転者については、改善基準告示に定める運転時間の限度を超えて運転業務に従事させることはできない。",
        size=FONT_SIZE_ARTICLE,
    )

    # 第8条（通知）
    add_paragraph(
        doc,
        "第８条　甲は、時間外労働を行わせる場合は、原則として、前日の終業時刻までに該当労働者に通知する。"
        "また、休日労働を行わせる場合は、原則として、１日前の終業時刻までに該当労働者に通知する。",
        size=FONT_SIZE_ARTICLE,
    )

    # 第9条（起算日・有効期間）
    add_paragraph(
        doc,
        f"第９条　第２条及び第４条の表における１年の起算日はいずれも{起算日}とする。",
        size=FONT_SIZE_ARTICLE,
    )
    add_paragraph(
        doc,
        f"　２　本協定の有効期間は、{起算日}から１年間とする。",
        size=FONT_SIZE_ARTICLE,
    )

    # 第2章（第10〜17条）
    _generate_chapter2_1nen(doc, data, art_start=10, 起算日=起算日)

    # 署名欄
    _generate_footer(doc, data)


# ---------------------------------------------------------------------------
# 様式パターン → 生成関数のマッピング
# ---------------------------------------------------------------------------
GENERATORS: Dict[str, Any] = {
    "9":    generate_form_9,
    "9_2":  generate_form_9_2,
    "9_3":  generate_form_9_3,
    "9_4":  generate_form_9_4,
    "9_5":  generate_form_9_5,
    "10":   generate_form_10,
    "10_2": generate_form_10_2,
}

FORM_NAMES: Dict[str, str] = {
    "9":    "様式第9号（一般条項）",
    "9_2":  "様式第9号の2（特別条項付き）",
    "9_3":  "様式第9号の3（研究開発業務）",
    "9_4":  "様式第9号の4（適用猶予事業）",
    "9_5":  "様式第9号の5（適用猶予＋事業場外みなし）",
    "10":   "様式第10号（1年変形制）",
    "10_2": "様式第10号の2（1年変形制＋ドライバー）",
}


def generate_word(data: Dict[str, Any], output_dir: str = "output") -> str:
    """Excelデータ1行分からWord協定書を生成する

    Args:
        data: 入力データ辞書
        output_dir: 出力ディレクトリパス

    Returns:
        生成されたファイルパス

    Raises:
        ValueError: Word生成またはファイル保存に失敗した場合
    """
    # _v() を使い、keyがあってもNoneの場合にデフォルト値を返す
    form_type: str = _v(data, "様式パターン") or "9"
    社名_raw: str = _v(data, "事業所名", "不明")

    generator = GENERATORS.get(form_type, generate_form_9)
    if form_type not in GENERATORS:
        logger.warning("未知の様式パターン '%s' を標準様式(9)で処理します [%s]", form_type, 社名_raw)
    form_name: str = FORM_NAMES.get(form_type, FORM_NAMES["9"])

    try:
        doc = Document()

        # ページ設定
        section = doc.sections[0]
        section.page_width = Mm(PAGE_WIDTH_MM)
        section.page_height = Mm(PAGE_HEIGHT_MM)
        section.top_margin = Cm(MARGIN_CM)
        section.bottom_margin = Cm(MARGIN_CM)
        section.left_margin = Cm(MARGIN_CM)
        section.right_margin = Cm(MARGIN_CM)

        # 文書生成
        logger.info("様式 %s の協定書を生成開始: %s", form_type, 社名_raw)
        generator(doc, data)

        # 保存
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)

        # ファイル名の安全化（/ \\ : * ? " < > | & % # をアンダースコアに置換）
        safe_name: str = _safe_filename(社名_raw)
        filename: str = f"36協定書_{safe_name}_{form_name}.docx"
        filepath: Path = output_path / filename

        doc.save(str(filepath))
        logger.info("生成完了: %s", filepath)
        return str(filepath)

    except Exception as exc:
        logger.error("Word生成エラー [様式=%s, 事業所=%s]: %s", form_type, 社名_raw, exc, exc_info=True)
        raise ValueError(f"Word生成に失敗しました（{社名_raw}、様式{form_type}）: {exc}") from exc


if __name__ == "__main__":
    # ログ設定（スクリプト直接実行時）
    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

    # テスト: サンプルデータで生成
    sample: Dict[str, Any] = {
        "事業所名": "テスト株式会社",
        "事業主名": "山田太郎",
        "電話番号": "03-1234-5678",
        "事業の種類": "情報通信業",
        "時間外_事由": "受注の集中、納期の逼迫",
        "時間外_業務の種類": "システム開発",
        "労働者数": "10",
        "所定労働時間": "8",
        "延長時間_1日": "4",
        "延長時間_1ヶ月": "45",
        "休日_事由": "納期対応",
        "休日_業務の種類": "システム開発",
        "所定休日": "土曜日・日曜日",
        "休日労働_日数": "2",
        "始業終業時刻": "9:00〜18:00",
        "起算日_年": "8",
        "起算日_月": "4",
        "様式パターン": "9",
        "労働者代表_職": "主任",
        "労働者代表_氏名": "鈴木一郎",
        "事業主職名": "代表取締役",
        "所轄労働局": "東京",
        "所轄労基署": "品川",
    }

    path = generate_word(sample, "output")
    logger.info("生成完了: %s", path)
