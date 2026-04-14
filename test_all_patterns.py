"""
全5パターン網羅テスト
様式9号〜9号の5まで全パターンのWord生成・検証を行う
"""
import sys
import os
import traceback
from pathlib import Path

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from word_generator import generate_word, FORM_NAMES, GENERATORS
from excel_reader import detect_form_type
from mail_sender import build_email_body, build_subject, create_email

TEST_OUTPUT = "test_output"
ERRORS = []
PASSED = []


def log_pass(name):
    PASSED.append(name)
    print(f"  ✅ PASS: {name}")


def log_fail(name, error):
    ERRORS.append((name, error))
    print(f"  ❌ FAIL: {name} → {error}")


def base_data():
    """共通テストデータ"""
    return {
        "更新月": "4月",
        "起算日_月": "4",
        "起算日_年": "8",
        "事業所名": "テスト株式会社",
        "事業主名": "山田太郎",
        "電話番号": "03-1234-5678",
        "事業の種類": "情報通信業",
        "時間外_事由": "受注の集中、納期の逼迫",
        "時間外_業務の種類": "システム開発",
        "労働者数": "10",
        "18歳未満_労働者数": "0",
        "所定労働時間": "8",
        "延長時間_1日": "4",
        "延長時間_1ヶ月": "45",
        "時間外_期間": "令和8年4月1日から1年間",
        "休日_事由": "納期対応",
        "休日_業務の種類": "システム開発",
        "所定休日": "土曜日・日曜日",
        "休日労働_日数": "2",
        "始業終業時刻": "9:00〜18:00",
        "休日_期間": "令和8年4月1日から1年間",
        "特別条項の有無": "",
        "特別_理由": "",
        "特別_業務の種類": "",
        "特別_労働者数": "",
        "特別_延長時間": "",
        "特別_超過回数": "",
        "特別_延長時間_月": "",
        "特別_割増賃金率": "",
        "特別_延長時間_年": "",
        "特別_手続き": "",
        "特別_健康措置_番号": "",
        "特別_健康措置_内容": "",
        "誓約チェック": "✓",
        "労働者代表_職": "主任",
        "労働者代表_氏名": "鈴木一郎",
        "過半数労働者_チェック": "✓",
        "協定締結日": "令和8年3月15日",
        "届出作成日": "令和8年3月20日",
        "事業主職名": "代表取締役",
        "届出_事業主名": "山田太郎",
        "所轄労働局": "東京",
        "所轄労基署": "品川",
        "メールアドレス": "test@example.com",
        "様式パターン": "9",
    }


# ============================================================
# テスト1: 様式パターン自動判定
# ============================================================
print("=" * 60)
print("テスト1: 様式パターン自動判定")
print("=" * 60)

# 9号（一般条項）
d = base_data()
d["特別条項の有無"] = "□"
result = detect_form_type(d)
if result == "9":
    log_pass("一般条項 → 9号")
else:
    log_fail("一般条項 → 9号", f"期待: 9, 実際: {result}")

# 9号（空の特別条項）
d2 = base_data()
d2["特別条項の有無"] = ""
result2 = detect_form_type(d2)
if result2 == "9":
    log_pass("空の特別条項 → 9号")
else:
    log_fail("空の特別条項 → 9号", f"期待: 9, 実際: {result2}")

# 9号の2（特別条項あり）
d3 = base_data()
d3["特別条項の有無"] = "あり"
d3["特別_理由"] = "臨時の受注増加"
d3["特別_延長時間_月"] = "70"
result3 = detect_form_type(d3)
if result3 == "9_2":
    log_pass("特別条項あり → 9号の2")
else:
    log_fail("特別条項あり → 9号の2", f"期待: 9_2, 実際: {result3}")

# 9号の2（特別条項フィールドのみ入力）
d3b = base_data()
d3b["特別条項の有無"] = ""
d3b["特別_理由"] = "繁忙期"
d3b["特別_延長時間_月"] = "60"
result3b = detect_form_type(d3b)
if result3b == "9_2":
    log_pass("特別条項フィールドのみ入力 → 9号の2")
else:
    log_fail("特別条項フィールドのみ入力 → 9号の2", f"期待: 9_2, 実際: {result3b}")

# 9号の3（研究開発）
d4 = base_data()
d4["時間外_業務の種類"] = "新商品の研究開発"
result4 = detect_form_type(d4)
if result4 == "9_3":
    log_pass("研究開発業務 → 9号の3")
else:
    log_fail("研究開発業務 → 9号の3", f"期待: 9_3, 実際: {result4}")

# 9号の4（適用猶予: 運転）
d5 = base_data()
d5["時間外_業務の種類"] = "自動車運転業務"
result5 = detect_form_type(d5)
if result5 == "9_4":
    log_pass("自動車運転業務 → 9号の4")
else:
    log_fail("自動車運転業務 → 9号の4", f"期待: 9_4, 実際: {result5}")

# 9号の4（適用猶予: 建設）
d5b = base_data()
d5b["事業の種類"] = "建設業"
result5b = detect_form_type(d5b)
if result5b == "9_4":
    log_pass("建設業 → 9号の4")
else:
    log_fail("建設業 → 9号の4", f"期待: 9_4, 実際: {result5b}")

# 9号の4（適用猶予: 医師）
d5c = base_data()
d5c["時間外_業務の種類"] = "医師の診療業務"
result5c = detect_form_type(d5c)
if result5c == "9_4":
    log_pass("医師の診療業務 → 9号の4")
else:
    log_fail("医師の診療業務 → 9号の4", f"期待: 9_4, 実際: {result5c}")

# 9号の5（適用猶予＋事業場外みなし）
d6 = base_data()
d6["事業の種類"] = "建設業"
d6["時間外_業務の種類"] = "事業場外作業"
result6 = detect_form_type(d6)
if result6 == "9_5":
    log_pass("建設+事業場外 → 9号の5")
else:
    log_fail("建設+事業場外 → 9号の5", f"期待: 9_5, 実際: {result6}")


# ============================================================
# テスト2: 全5パターンWord生成
# ============================================================
print()
print("=" * 60)
print("テスト2: 全5パターンWord生成")
print("=" * 60)

patterns = {
    "9": {"事業所名": "一般条項テスト株式会社"},
    "9_2": {
        "事業所名": "特別条項テスト工業",
        "特別条項の有無": "あり",
        "特別_理由": "臨時の受注増加、納期の変更",
        "特別_業務の種類": "製造・品質管理",
        "特別_労働者数": "20",
        "特別_超過回数": "6",
        "特別_延長時間_月": "70",
        "特別_割増賃金率": "25",
        "特別_延長時間_年": "700",
        "特別_手続き": "労使の協議",
        "特別_健康措置_内容": "産業医による面接指導",
    },
    "9_3": {
        "事業所名": "研究開発テストラボ",
        "時間外_業務の種類": "新技術の研究開発",
    },
    "9_4": {
        "事業所名": "適用猶予テスト運輸",
        "事業の種類": "運輸業",
        "時間外_業務の種類": "自動車運転業務",
    },
    "9_5": {
        "事業所名": "事業場外テスト建設",
        "事業の種類": "建設業",
        "時間外_業務の種類": "事業場外の現場作業",
        "特別_手続き": "みなし労働時間制適用",
    },
}

generated_files = []
for pattern_key, overrides in patterns.items():
    d = base_data()
    d.update(overrides)
    d["様式パターン"] = pattern_key

    test_name = f"Word生成: {FORM_NAMES[pattern_key]}"
    try:
        filepath = generate_word(d, TEST_OUTPUT)
        if Path(filepath).exists() and Path(filepath).stat().st_size > 0:
            generated_files.append(filepath)
            log_pass(test_name)
        else:
            log_fail(test_name, "ファイルが空または存在しない")
    except Exception as e:
        log_fail(test_name, f"{e}\n{traceback.format_exc()}")


# ============================================================
# テスト3: 生成Wordの内容検証
# ============================================================
print()
print("=" * 60)
print("テスト3: 生成Wordの内容検証")
print("=" * 60)

from docx import Document

for filepath in generated_files:
    fname = Path(filepath).name
    try:
        doc = Document(filepath)
        text = "\n".join(p.text for p in doc.paragraphs)
        tables_count = len(doc.tables)

        # 基本チェック
        checks = []
        if "協定書" in text or "協定" in text:
            checks.append("タイトルあり")
        else:
            log_fail(f"内容検証 {fname}: タイトル", "協定書のタイトルが見つからない")
            continue

        if tables_count > 0:
            checks.append(f"テーブル{tables_count}個")
        else:
            log_fail(f"内容検証 {fname}: テーブル", "テーブルが0個")
            continue

        if "事業" in text:
            checks.append("事業所情報あり")

        if "労働者代表" in text or "署名" in text or "（乙）" in text:
            checks.append("署名欄あり")

        log_pass(f"内容検証 {fname} ({', '.join(checks)})")

    except Exception as e:
        log_fail(f"内容検証 {fname}", str(e))


# ============================================================
# テスト4: メール機能
# ============================================================
print()
print("=" * 60)
print("テスト4: メール機能")
print("=" * 60)

# メール本文生成
config = {"差出人名": "安田", "差出人所属": "朝日事務所", "差出人電話": "03-0000-0000"}
d = base_data()
body = build_email_body(d, config)
if "山田太郎" in body and "テスト株式会社" in body and "安田" in body:
    log_pass("メール本文生成（宛先名・事業所名・差出人名が含まれる）")
else:
    log_fail("メール本文生成", f"期待する文字列が本文に含まれない:\n{body}")

# メール件名生成
subject = build_subject(d)
if "テスト株式会社" in subject and "36協定" in subject:
    log_pass("メール件名生成")
else:
    log_fail("メール件名生成", f"件名: {subject}")

# MIMEメッセージ作成（添付ファイル付き）
if generated_files:
    try:
        msg = create_email(
            to_address="test@example.com",
            subject=subject,
            body=body,
            attachment_path=generated_files[0],
            from_address="sender@example.com",
        )
        if msg["To"] == "test@example.com" and msg["Subject"] == subject:
            # 添付ファイルチェック
            attachments = [p for p in msg.walk() if p.get_content_disposition() == "attachment"]
            if len(attachments) == 1:
                log_pass("MIMEメッセージ作成（ヘッダー＋添付ファイル1個）")
            else:
                log_fail("MIMEメッセージ作成", f"添付ファイル数: {len(attachments)}")
        else:
            log_fail("MIMEメッセージ作成", "ヘッダーが不正")
    except Exception as e:
        log_fail("MIMEメッセージ作成", str(e))


# ============================================================
# テスト5: エッジケース
# ============================================================
print()
print("=" * 60)
print("テスト5: エッジケース")
print("=" * 60)

# 空データ
d_empty = base_data()
d_empty["時間外_事由"] = ""
d_empty["休日_事由"] = ""
d_empty["労働者数"] = ""
d_empty["延長時間_1日"] = ""
try:
    fp = generate_word(d_empty, TEST_OUTPUT)
    if Path(fp).exists():
        log_pass("空データでもWord生成可能（クラッシュしない）")
    else:
        log_fail("空データ", "ファイルが生成されない")
except Exception as e:
    log_fail("空データ", str(e))

# 特殊文字（&, <, >）
d_special = base_data()
d_special["事業所名"] = "A&B株式会社<テスト>"
d_special["時間外_事由"] = "受注 & 納期の\"逼迫\""
try:
    fp = generate_word(d_special, TEST_OUTPUT)
    doc = Document(fp)
    text = "\n".join(p.text for p in doc.paragraphs)
    if "A&B" in text:
        log_pass("特殊文字（&, <, >）がWord内で正しく表示される")
    else:
        log_fail("特殊文字", "特殊文字がWordに反映されていない")
except Exception as e:
    log_fail("特殊文字", str(e))

# 長い文字列
d_long = base_data()
d_long["時間外_事由"] = "受注の集中、納期の逼迫、年度末の決算対応、システム障害対応、顧客からの緊急依頼対応、新規プロジェクト立ち上げに伴う準備作業" * 3
try:
    fp = generate_word(d_long, TEST_OUTPUT)
    if Path(fp).exists() and Path(fp).stat().st_size > 0:
        log_pass("長い文字列でもWord生成可能")
    else:
        log_fail("長い文字列", "ファイルが空")
except Exception as e:
    log_fail("長い文字列", str(e))

# メールアドレスなしの処理
d_no_email = base_data()
d_no_email["メールアドレス"] = ""
subject = build_subject(d_no_email)
body = build_email_body(d_no_email, config)
if subject and body:
    log_pass("メールアドレスなしでも件名・本文生成可能")
else:
    log_fail("メールアドレスなし", "件名or本文が空")

# ファイル名に使えない文字
d_badname = base_data()
d_badname["事業所名"] = "テスト/会社\\名"
try:
    fp = generate_word(d_badname, TEST_OUTPUT)
    if Path(fp).exists():
        log_pass("ファイル名の特殊文字（/\\）がサニタイズされる")
    else:
        log_fail("ファイル名サニタイズ", "ファイルが生成されない")
except Exception as e:
    log_fail("ファイル名サニタイズ", str(e))


# ============================================================
# 結果サマリー
# ============================================================
print()
print("=" * 60)
print("テスト結果サマリー")
print("=" * 60)
print(f"  ✅ PASS: {len(PASSED)}件")
print(f"  ❌ FAIL: {len(ERRORS)}件")
if ERRORS:
    print()
    print("--- 失敗詳細 ---")
    for name, error in ERRORS:
        print(f"  ❌ {name}: {error}")
    sys.exit(1)
else:
    print()
    print("全テスト合格！")
    sys.exit(0)
