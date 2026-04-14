"""
36協定自動化 pytest テストスイート
excel_reader, mail_sender, word_generator の各モジュールを網羅的にテストする
"""
import os
import sys
import tempfile
from pathlib import Path
from unittest.mock import patch, MagicMock

import pytest
import openpyxl

# テスト対象モジュールのインポート
sys.path.insert(0, str(Path(__file__).parent))
from excel_reader import (
    detect_form_type,
    validate_record,
    read_excel,
    COLUMN_MAP,
    EMAIL_COLUMN,
    FORM_TYPE_RULES,
    DEFAULT_FORM_TYPE,
)
from mail_sender import build_email_body, build_subject, create_email, DEFAULT_TEMPLATE
from word_generator import generate_word, FORM_NAMES, GENERATORS


# ====================================================================
# ヘルパー: テスト用レコード・Excelファイル生成
# ====================================================================

def _base_record(**overrides) -> dict:
    """基本レコードを生成し、overridesで上書きする"""
    record = {
        "事業所名": "テスト株式会社",
        "事業主名": "山田太郎",
        "電話番号": "03-1234-5678",
        "事業の種類": "情報通信業",
        "時間外_事由": "受注の集中",
        "時間外_業務の種類": "システム開発",
        "労働者数": "10",
        "18歳未満_労働者数": "0",
        "所定労働時間": "8",
        "延長時間_1日": "4",
        "延長時間_1ヶ月": "45",
        "時間外_期間": "",
        "休日_事由": "納期対応",
        "休日_業務の種類": "システム開発",
        "所定休日": "土曜日・日曜日",
        "休日労働_日数": "2",
        "始業終業時刻": "9:00〜18:00",
        "休日_期間": "",
        "特別条項の有無": "",
        "特別_理由": "",
        "特別_業務の種類": "",
        "特別_労働者数": "",
        "特別_延長時間": "",
        "特別_超過回数": "",
        "特別_延長時間_月": "",
        "特別_割増賃金率": "",
        "特別_延長時間_年": "",
        "特別_手続き": "",
        "特別_健康措置_番号": "",
        "特別_健康措置_内容": "",
        "誓約チェック": "",
        "労働者代表_職": "主任",
        "労働者代表_氏名": "鈴木一郎",
        "過半数労働者_チェック": "",
        "協定締結日": "2026年4月1日",
        "届出作成日": "2026年3月20日",
        "事業主職名": "代表取締役",
        "届出_事業主名": "山田太郎",
        "所轄労働局": "東京",
        "所轄労基署": "品川",
        "更新月": "4",
        "起算日_月": "4",
        "起算日_年": "8",
        "メールアドレス": "test@example.com",
    }
    record.update(overrides)
    return record


def _create_test_excel(records: list[dict], path: str) -> str:
    """テスト用Excelファイルを生成する"""
    wb = openpyxl.Workbook()
    ws = wb.active

    # ヘッダー行（1行目）— 実際は使わないがread_excelが2行目からなので
    ws.cell(row=1, column=1, value="ヘッダー")

    for row_idx, record in enumerate(records, start=2):
        for key, col in COLUMN_MAP.items():
            val = record.get(key, "")
            if val:
                ws.cell(row=row_idx, column=col, value=val)
        # メールアドレス
        email = record.get("メールアドレス", "")
        if email:
            ws.cell(row=row_idx, column=EMAIL_COLUMN, value=email)

    wb.save(path)
    wb.close()
    return path


# ====================================================================
# TestDetectFormType: 全5パターン + エッジケース（15テスト以上）
# ====================================================================

class TestDetectFormType:
    """detect_form_type() の様式判定テスト"""

    # --- 基本5パターン ---

    def test_default_form_9(self):
        """一般条項（デフォルト）→ 様式第9号"""
        record = _base_record()
        assert detect_form_type(record) == "9"

    def test_form_9_2_tokubetsu_flag(self):
        """特別条項フラグあり → 様式第9号の2"""
        record = _base_record(特別条項の有無="あり")
        assert detect_form_type(record) == "9_2"

    def test_form_9_3_kenkyuu(self):
        """研究開発業務 → 様式第9号の3"""
        record = _base_record(時間外_業務の種類="新技術の研究開発")
        assert detect_form_type(record) == "9_3"

    def test_form_9_4_yuuyo_unten(self):
        """適用猶予（運転）→ 様式第9号の4"""
        record = _base_record(事業の種類="自動車運転業")
        assert detect_form_type(record) == "9_4"

    def test_form_9_5_yuuyo_minashi(self):
        """適用猶予＋みなし → 様式第9号の5"""
        record = _base_record(事業の種類="建設業", 特別_手続き="みなし労働")
        assert detect_form_type(record) == "9_5"

    # --- エッジケース ---

    def test_form_9_2_by_reason(self):
        """特別条項の有無は空だが特別_理由がある → 9_2"""
        record = _base_record(特別_理由="臨時的な受注増")
        assert detect_form_type(record) == "9_2"

    def test_form_9_2_by_monthly_extension(self):
        """特別条項の有無は空だが特別_延長時間_月がある → 9_2"""
        record = _base_record(特別_延長時間_月="80")
        assert detect_form_type(record) == "9_2"

    def test_tokubetsu_nashi_remains_9(self):
        """特別条項の有無が「なし」→ デフォルト9"""
        record = _base_record(特別条項の有無="なし")
        assert detect_form_type(record) == "9"

    def test_tokubetsu_checkbox_remains_9(self):
        """特別条項の有無が「□」→ デフォルト9"""
        record = _base_record(特別条項の有無="□")
        assert detect_form_type(record) == "9"

    def test_yuuyo_kensetsu(self):
        """建設業 → 適用猶予 9_4"""
        record = _base_record(事業の種類="建設業")
        assert detect_form_type(record) == "9_4"

    def test_yuuyo_ishi(self):
        """医師 → 適用猶予 9_4"""
        record = _base_record(時間外_業務の種類="医師の診療業務")
        assert detect_form_type(record) == "9_4"

    def test_yuuyo_kagoshima(self):
        """鹿児島・砂糖 → 適用猶予 9_4"""
        record = _base_record(事業の種類="鹿児島の砂糖製造")
        assert detect_form_type(record) == "9_4"

    def test_yuuyo_okinawa(self):
        """沖縄 → 適用猶予 9_4"""
        record = _base_record(事業の種類="沖縄の砂糖製造")
        assert detect_form_type(record) == "9_4"

    def test_yuuyo_takes_priority_over_kenkyuu(self):
        """猶予キーワードと研究キーワードが両方ある場合、猶予が優先"""
        record = _base_record(事業の種類="建設業", 時間外_業務の種類="新技術の研究")
        # 猶予ルール（9_4）が研究ルール（9_3）より先に評価される
        assert detect_form_type(record) == "9_4"

    def test_yuuyo_minashi_jigyoubagai(self):
        """事業場外キーワードでみなし判定 → 9_5"""
        record = _base_record(事業の種類="運転業", 時間外_業務の種類="事業場外営業")
        assert detect_form_type(record) == "9_5"

    def test_all_empty_returns_default(self):
        """全フィールド空 → デフォルト9"""
        record = {k: "" for k in COLUMN_MAP}
        assert detect_form_type(record) == "9"

    def test_shinshounin_kenkyuu(self):
        """新商品 → 研究開発 9_3"""
        record = _base_record(時間外_業務の種類="新商品の開発")
        assert detect_form_type(record) == "9_3"

    def test_rules_list_has_4_entries(self):
        """ルールリストが4エントリであることを確認"""
        assert len(FORM_TYPE_RULES) == 4

    def test_default_form_type_is_9(self):
        """デフォルト様式コードが "9" であること"""
        assert DEFAULT_FORM_TYPE == "9"


# ====================================================================
# TestValidateRecord: 必須項目チェック、数値チェック、月45h超（8テスト以上）
# ====================================================================

class TestValidateRecord:
    """validate_record() のバリデーションテスト"""

    def test_valid_record_no_warnings(self):
        """正常レコード → 警告なし"""
        record = _base_record()
        warnings = validate_record(record, 2)
        assert warnings == []

    def test_missing_jigyoushumei(self):
        """事業主名が空 → 必須項目警告"""
        record = _base_record(事業主名="")
        warnings = validate_record(record, 2)
        assert any("事業主名" in w for w in warnings)

    def test_missing_jigyou_shurui(self):
        """事業の種類が空 → 必須項目警告"""
        record = _base_record(事業の種類="")
        warnings = validate_record(record, 2)
        assert any("事業の種類" in w for w in warnings)

    def test_missing_gyoumu_shurui(self):
        """時間外_業務の種類が空 → 必須項目警告"""
        record = _base_record(時間外_業務の種類="")
        warnings = validate_record(record, 2)
        assert any("時間外_業務の種類" in w for w in warnings)

    def test_missing_roudousha_suu(self):
        """労働者数が空 → 必須項目警告"""
        record = _base_record(労働者数="")
        warnings = validate_record(record, 2)
        assert any("労働者数" in w for w in warnings)

    def test_multiple_missing_fields(self):
        """複数の必須項目が空 → 複数警告"""
        record = _base_record(事業主名="", 労働者数="")
        warnings = validate_record(record, 2)
        assert len(warnings) >= 2

    def test_non_numeric_roudousha(self):
        """労働者数が非数値 → 数値警告"""
        record = _base_record(労働者数="abc")
        warnings = validate_record(record, 2)
        assert any("数値ではありません" in w and "労働者数" in w for w in warnings)

    def test_non_numeric_enchou_ichinichi(self):
        """延長時間_1日が非数値 → 数値警告"""
        record = _base_record(延長時間_1日="三時間")
        warnings = validate_record(record, 2)
        assert any("延長時間_1日" in w for w in warnings)

    def test_over_45h_without_tokubetsu(self):
        """月45時間超で特別条項なし → 警告"""
        record = _base_record(延長時間_1ヶ月="60", 特別条項の有無="")
        warnings = validate_record(record, 2)
        assert any("45時間超" in w for w in warnings)

    def test_over_45h_with_tokubetsu(self):
        """月45時間超で特別条項あり → 警告なし"""
        record = _base_record(延長時間_1ヶ月="60", 特別条項の有無="あり")
        warnings = validate_record(record, 2)
        assert not any("45時間超" in w for w in warnings)

    def test_exactly_45h_no_warning(self):
        """月ちょうど45時間 → 45時間超警告なし"""
        record = _base_record(延長時間_1ヶ月="45", 特別条項の有無="")
        warnings = validate_record(record, 2)
        assert not any("45時間超" in w for w in warnings)

    def test_row_number_in_warning(self):
        """警告メッセージに行番号が含まれる"""
        record = _base_record(事業主名="")
        warnings = validate_record(record, 5)
        assert any("行5" in w for w in warnings)

    def test_jigyousho_name_in_warning(self):
        """警告メッセージに事業所名が含まれる"""
        record = _base_record(事業所名="サンプル社", 事業主名="")
        warnings = validate_record(record, 2)
        assert any("サンプル社" in w for w in warnings)

    def test_decimal_is_valid_number(self):
        """小数（例: 3.5）は数値として有効"""
        record = _base_record(延長時間_1日="3.5")
        warnings = validate_record(record, 2)
        assert not any("延長時間_1日" in w and "数値ではありません" in w for w in warnings)


# ====================================================================
# TestReadExcel: 正常系、空ファイル、不正形式（5テスト以上）
# ====================================================================

class TestReadExcel:
    """read_excel() のExcel読み取りテスト"""

    def test_read_single_record(self, tmp_path):
        """1レコードのExcelを正しく読み取れる"""
        record = _base_record()
        xlsx_path = str(tmp_path / "test.xlsx")
        _create_test_excel([record], xlsx_path)

        results = read_excel(xlsx_path)
        assert len(results) == 1
        assert results[0]["事業所名"] == "テスト株式会社"

    def test_read_multiple_records(self, tmp_path):
        """複数レコードのExcelを正しく読み取れる"""
        records = [
            _base_record(事業所名="A社"),
            _base_record(事業所名="B社"),
            _base_record(事業所名="C社"),
        ]
        xlsx_path = str(tmp_path / "test.xlsx")
        _create_test_excel(records, xlsx_path)

        results = read_excel(xlsx_path)
        assert len(results) == 3
        names = [r["事業所名"] for r in results]
        assert "A社" in names
        assert "B社" in names
        assert "C社" in names

    def test_read_empty_excel(self, tmp_path):
        """データ行がないExcel → 空リスト"""
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.cell(row=1, column=1, value="ヘッダー")
        xlsx_path = str(tmp_path / "empty.xlsx")
        wb.save(xlsx_path)
        wb.close()

        results = read_excel(xlsx_path)
        assert results == []

    def test_skip_row_without_jigyousho(self, tmp_path):
        """事業所名が空の行はスキップされる"""
        records = [
            _base_record(事業所名="A社"),
            _base_record(事業所名=""),  # スキップされるべき
            _base_record(事業所名="C社"),
        ]
        xlsx_path = str(tmp_path / "test.xlsx")
        _create_test_excel(records, xlsx_path)

        results = read_excel(xlsx_path)
        assert len(results) == 2
        names = [r["事業所名"] for r in results]
        assert "A社" in names
        assert "C社" in names

    def test_form_type_auto_detected(self, tmp_path):
        """読み取り時に様式パターンが自動付与される"""
        records = [_base_record(特別条項の有無="あり")]
        xlsx_path = str(tmp_path / "test.xlsx")
        _create_test_excel(records, xlsx_path)

        results = read_excel(xlsx_path)
        assert results[0]["様式パターン"] == "9_2"

    def test_email_column_read(self, tmp_path):
        """メールアドレスが正しく読み取れる"""
        records = [_base_record(メールアドレス="user@example.com")]
        xlsx_path = str(tmp_path / "test.xlsx")
        _create_test_excel(records, xlsx_path)

        results = read_excel(xlsx_path)
        assert results[0]["メールアドレス"] == "user@example.com"

    def test_nonexistent_file_raises(self):
        """存在しないファイルを開くと例外が発生する"""
        with pytest.raises(Exception):
            read_excel("/tmp/nonexistent_file_12345.xlsx")

    def test_invalid_format_raises(self, tmp_path):
        """不正なファイル形式（テキスト）は例外を発生させる"""
        bad_path = str(tmp_path / "bad.xlsx")
        with open(bad_path, "w") as f:
            f.write("これはExcelではない")
        with pytest.raises(Exception):
            read_excel(bad_path)


# ====================================================================
# TestBuildEmailBody: 正常系、空データ（3テスト以上）
# ====================================================================

class TestBuildEmailBody:
    """build_email_body() のメール本文生成テスト"""

    def test_normal_body(self):
        """正常系: テンプレートにデータが埋め込まれる"""
        data = {"事業主名": "佐藤花子", "事業所名": "サクラ株式会社"}
        config = {"差出人名": "安田", "差出人所属": "朝日事務所", "差出人電話": "03-0000-0000"}

        body = build_email_body(data, config)
        assert "佐藤花子" in body
        assert "サクラ株式会社" in body
        assert "安田" in body
        assert "朝日事務所" in body

    def test_empty_data_uses_defaults(self):
        """空データでもデフォルト値が使われてエラーにならない"""
        data = {}
        config = {"差出人名": "", "差出人所属": "", "差出人電話": ""}

        body = build_email_body(data, config)
        assert "ご担当者" in body  # デフォルトの宛先名

    def test_custom_template(self):
        """カスタムテンプレートが使われる"""
        data = {"事業主名": "田中太郎", "事業所名": "タナカ社"}
        config = {
            "メールテンプレート": "{宛先名}様へ。{事業主名}の書類です。{差出人名}{差出人所属}{差出人電話}",
            "差出人名": "安田",
            "差出人所属": "事務所",
            "差出人電話": "000",
        }
        body = build_email_body(data, config)
        # 宛先名はdata["事業主名"]、事業主名はdata["事業所名"]が埋め込まれる仕様
        assert "田中太郎様へ" in body
        assert "タナカ社の書類です" in body
        assert "安田" in body

    def test_body_contains_signature(self):
        """デフォルトテンプレートに署名区切り線が含まれる"""
        data = {"事業主名": "テスト", "事業所名": "テスト社"}
        config = {"差出人名": "A", "差出人所属": "B", "差出人電話": "C"}
        body = build_email_body(data, config)
        assert "──────" in body


# ====================================================================
# TestBuildSubject: 正常系（2テスト以上）
# ====================================================================

class TestBuildSubject:
    """build_subject() のメール件名生成テスト"""

    def test_normal_subject(self):
        """正常系: 事業所名が件名に含まれる"""
        data = {"事業所名": "サクラ株式会社"}
        subject = build_subject(data)
        assert "サクラ株式会社" in subject
        assert "36協定届" in subject

    def test_empty_name(self):
        """事業所名が空でも件名が生成される"""
        data = {"事業所名": ""}
        subject = build_subject(data)
        assert "36協定届" in subject

    def test_subject_format(self):
        """件名が「【36協定届】〇〇様 ...」の形式になる"""
        data = {"事業所名": "ABC社"}
        subject = build_subject(data)
        assert subject.startswith("【36協定届】")
        assert "ABC社様" in subject


# ====================================================================
# TestGenerateWord: 全5パターン生成＋内容検証（10テスト以上）
# ====================================================================

class TestGenerateWord:
    """generate_word() のWord生成テスト"""

    def test_generate_form_9(self, tmp_path):
        """様式第9号（一般条項）のWord生成"""
        data = _base_record(様式パターン="9")
        filepath = generate_word(data, str(tmp_path))
        assert Path(filepath).exists()
        assert "様式第9号" in filepath

    def test_generate_form_9_2(self, tmp_path):
        """様式第9号の2（特別条項付き）のWord生成"""
        data = _base_record(
            様式パターン="9_2",
            特別条項の有無="あり",
            特別_理由="臨時的な受注増",
            特別_業務の種類="システム開発",
            特別_労働者数="5",
            特別_超過回数="6",
            特別_延長時間_月="80",
            特別_延長時間_年="720",
            特別_割増賃金率="25",
            特別_健康措置_内容="医師面談",
        )
        filepath = generate_word(data, str(tmp_path))
        assert Path(filepath).exists()
        assert "特別条項" in filepath

    def test_generate_form_9_3(self, tmp_path):
        """様式第9号の3（研究開発業務）のWord生成"""
        data = _base_record(様式パターン="9_3", 時間外_業務の種類="新技術の研究開発")
        filepath = generate_word(data, str(tmp_path))
        assert Path(filepath).exists()
        assert "研究開発" in filepath

    def test_generate_form_9_4(self, tmp_path):
        """様式第9号の4（適用猶予事業）のWord生成"""
        data = _base_record(様式パターン="9_4", 事業の種類="建設業")
        filepath = generate_word(data, str(tmp_path))
        assert Path(filepath).exists()
        assert "適用猶予" in filepath

    def test_generate_form_9_5(self, tmp_path):
        """様式第9号の5（適用猶予＋事業場外みなし）のWord生成"""
        data = _base_record(様式パターン="9_5", 事業の種類="建設業")
        filepath = generate_word(data, str(tmp_path))
        assert Path(filepath).exists()
        assert "みなし" in filepath

    def test_output_dir_created(self, tmp_path):
        """出力ディレクトリが存在しない場合は自動作成される"""
        output_dir = str(tmp_path / "new_dir" / "sub")
        data = _base_record(様式パターン="9")
        filepath = generate_word(data, output_dir)
        assert Path(filepath).exists()
        assert Path(output_dir).is_dir()

    def test_filename_contains_jigyousho(self, tmp_path):
        """ファイル名に事業所名が含まれる"""
        data = _base_record(事業所名="富士通テスト", 様式パターン="9")
        filepath = generate_word(data, str(tmp_path))
        assert "富士通テスト" in Path(filepath).name

    def test_filename_safe_characters(self, tmp_path):
        """事業所名にスラッシュ・バックスラッシュが含まれても安全なファイル名になる"""
        data = _base_record(事業所名="A/B\\C 社", 様式パターン="9")
        filepath = generate_word(data, str(tmp_path))
        assert Path(filepath).exists()
        name = Path(filepath).name
        assert "/" not in name
        assert "\\" not in name

    def test_docx_extension(self, tmp_path):
        """生成ファイルが.docx拡張子を持つ"""
        data = _base_record(様式パターン="9")
        filepath = generate_word(data, str(tmp_path))
        assert filepath.endswith(".docx")

    def test_all_7_generators_exist(self):
        """GENERATORSに全7パターンが登録されている（10/10_2を含む）"""
        expected = {"9", "9_2", "9_3", "9_4", "9_5", "10", "10_2"}
        assert set(GENERATORS.keys()) == expected

    def test_all_7_form_names_exist(self):
        """FORM_NAMESに全7パターンが登録されている（10/10_2を含む）"""
        expected = {"9", "9_2", "9_3", "9_4", "9_5", "10", "10_2"}
        assert set(FORM_NAMES.keys()) == expected

    def test_word_content_contains_jigyousho(self, tmp_path):
        """生成されたWordに事業所名が含まれる"""
        from docx import Document
        data = _base_record(事業所名="コンテンツ検証社", 様式パターン="9")
        filepath = generate_word(data, str(tmp_path))
        doc = Document(filepath)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "コンテンツ検証社" in full_text

    def test_word_content_contains_jigyoushumei(self, tmp_path):
        """生成されたWordに事業主名が含まれる"""
        from docx import Document
        data = _base_record(事業主名="検証太郎", 様式パターン="9")
        filepath = generate_word(data, str(tmp_path))
        doc = Document(filepath)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "検証太郎" in full_text

    def test_form_9_2_contains_tokubetsu_clause(self, tmp_path):
        """様式9_2のWordに特別条項の内容（健康確保措置・割増率）が含まれる
        ※段落形式（第7条）のためヘッダー「特別条項」という文字列ではなく内容で検証"""
        from docx import Document
        data = _base_record(様式パターン="9_2", 特別条項の有無="あり", 特別_理由="テスト理由")
        filepath = generate_word(data, str(tmp_path))
        doc = Document(filepath)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        # 特別条項の実質内容（健康確保措置と割増率）が含まれること
        assert "健康確保措置" in full_text
        assert "割増率" in full_text
        # 特別条項は第7条として実装されている
        assert "第８条" in full_text  # 有効期間が第8条であることを確認


# ====================================================================
# 合計テスト数の確認（pytest -v で見える）
# ====================================================================

if __name__ == "__main__":
    pytest.main([__file__, "-v"])
